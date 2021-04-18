# -*- coding: utf-8 -*-
import docx
import fitz
import textract
import os
import re
from os import path
from PIL import Image
import io
import pytesseract
import cv2
import matplotlib.pyplot as plt
import collections
import re
from deeppavlov import configs, build_model, train_model
from deeppavlov.core.commands.utils import parse_config
from nltk.corpus import stopwords
from natasha import (
    Segmenter,
    MorphVocab,

    NewsEmbedding,
    NewsMorphTagger,
    NewsSyntaxParser,
    NewsNERTagger,

    PER,
    NamesExtractor,

    Doc
)

segmenter = Segmenter()
morph_vocab = MorphVocab()
emb = NewsEmbedding()
morph_tagger = NewsMorphTagger(emb)
syntax_parser = NewsSyntaxParser(emb)
ner_tagger = NewsNERTagger(emb)
names_extractor = NamesExtractor(morph_vocab)
ner_model = build_model(configs.ner.ner_rus_bert, download=False)


def get_feedback(document, image_name, mode):  # Получение(со скана) текста
    if mode == 'docx':  # Читаем скан с DOCX
        image2_info = []  # [rId, name]
        for r in document.part.rels.values():
            if isinstance(r._target, docx.parts.image.ImagePart) and ('{}.'.format(image_name) in r._target.partname):
                image2_info.append(r.rId)
                image2_info.append(os.path.basename(r._target.partname))

        doc_part = document.part
        image_part = doc_part.related_parts[image2_info[0]]
        a = image_part.blob

        image = Image.open(io.BytesIO(a))
        image.save(image2_info[1])
        img = image2_info[1]
    elif mode == 'pdf':  # Читаем скан с PDF
        for img in document.getPageImageList(0):
            xref = img[0]
            pix = fitz.Pixmap(document, xref)
            pix1 = fitz.Pixmap(fitz.csRGB, pix)
            img = '{}.png'.format(image_name)
            pix1.writePNG(img)
            pix1 = None
    k1 = cv2.imread(img)
    gray = cv2.cvtColor(k1, cv2.COLOR_BGR2GRAY)
    # threshold = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)[1]
    imn = image_name + '.jpeg'
    cv2.imwrite(imn, gray)

    text = pytesseract.image_to_string(Image.open(img), lang='rus')
    return text


def get_name_from_feedback1(feedback):
    try:
        name = ''
        res = ner_model([feedback])
        if res:
            for id, item in enumerate(res[1][0]):
                if item in ['I-PER', 'B-PER']:
                    name += ' ' + res[0][0][id]
            print('Ф.И.О: ', name)
            return name
        else:
            print('Ф.И.О: Error')
            return 'Error'
    except:
        print('Ф.И.О: Error')
        return 'Error'


def find_theme(text):
    try:
        pattern = r"(?<=на тему «).*?(?=»)"
        s_l = re.findall(pattern, text)
        if not s_l:
            pattern = r"(?<=на тему: «).*?(?=»)"
            s_l = re.findall(pattern, text)
            if not s_l:
                pattern = r"(?<=на тему — «).*?(?=»)"
                s_l = re.findall(pattern, text)
                if not s_l:
                    start_text = text.find('тема работы')
                    end_text = text.find('студ')
                    t_name = text[start_text + 11:end_text]
                    print('Тема:', t_name)
                    return t_name
        print('Тема:', s_l[0])
        return s_l[0]
    except:
        print('Тема: Error')
        return 'Error'


def find_profile(text):
    try:
        start_text = text.find('профиль')
        text = text[start_text:]
        pattern = r"(?<=«).*?(?=»)"
        s_l = re.findall(pattern, text)
        if s_l:
            print('Профиль:', s_l[0])
            return s_l[0]
        else:
            start_text = text.find('филь')
            end_text = text.find('студ')
            p_name = text[start_text + 4:end_text]
            print('Профиль:', p_name)
            return p_name
    except:
        print('Профиль: Error')
        return 'Error'


def find_direction(text):
    try:
        d_name = ''
        start_text = text.find('направление')
        pattern = r"(?<=«).*?(?=»)"
        s_l = re.findall(pattern, text[start_text:])
        if s_l:
            print('Направление:', s_l[0])
            return s_l[0]
        else:
            start_text = text.find('вление')
            end_text = text.find('проф')
            d_name = text[start_text + 6:end_text]
            if (start_text != -1 or start_text != -1) and d_name:
                print('Направление:', d_name)
                return d_name
            else:
                start_text = text.find('льность')
                end_text = text.find('студ')
                d_name = text[start_text + 7:end_text]
                print('Направление:', d_name)
                return d_name
    except:
        print('Направление: Error')
        return 'Error'


def find_faculty(text):
    try:
        pattern = r"(?<=факультет ).*?(?=кафедра)"
        a = re.findall(pattern, text)
        a = re.sub(r' {2,}', ' ', a[0])
        print('Факультет:', a)
        return a
    except:
        print('Факультет: Error')
        return 'Error'


def save_in_docx(data, dir, mode):  # сохраняем результат в Ворд файл
    name = os.path.basename(dir.split('.')[-1])
    file = '@{}{}.docx'.format(mode, name)
    if os.path.isfile(file):
        os.remove(file)
    doc_save = docx.Document()
    for item in data:
        doc_save.add_paragraph(item)
    doc_save.save(file)

def depart_find(text):
    try:
        start_text = text.find('кафедра')
        end_text = text.find('допус')
        dep_name = text[start_text + 7:end_text]
        return dep_name
    except:
        print('Кафедра: Error')
        return 'Error'


def load_docx_link(
        folder):  # загружаем все адреса word файлов из папки folder, исключаем временные файлы(~), и файлы начинающиеся на @
    paths = []
    for root, dirs, files in os.walk(folder):
        for file in files:
            if (file.name.split('.')[-1] == 'docx' or file.name.split('.')[-1] == 'doc' or file.name.split('.')[
                -1] == 'pdf') and not file.startswith('~') and not file.startswith('@'):
                paths.append(os.path.join(root, file))
    return paths


def text_or_scan(file):  # проверка документа на наличие скана в первой странице
    text = None
    if file.name.split('.')[-1] == 'doc':
        text = text_from_doc(file).splitlines()
        text = ' '.join(text).lower()
        print('это doc')
    elif file.name.split('.')[-1] == 'docx':
        print('это docx')
        doc = docx.Document(file)
        text = text_from_docx(doc).splitlines()
        text = ' '.join(text).lower()
    elif file.name.split('.')[-1] == 'pdf':
        print('Это PDF!')
        doc = fitz.open(file)
        page1 = doc.loadPage(0)
        text = page1.getText("text").splitlines()
        text = ' '.join(text).lower()
    if text:
        if ('ульяновский' in text and
            'факультет' in text and
            'кафедра' in text) or (

                'ульян' in text and
                'факул' in text and
                'каф' in text):
            # print('Текст')
            return True
        else:
            # print('Скан')
            return False


def most_common_word(text):  # частотный словарь
    try:
        word_cloud = '\n'
        doc = Doc(text)
        doc.segment(segmenter)
        doc.tag_morph(morph_tagger)
        doc.parse_syntax(syntax_parser)
        doc.tag_ner(ner_tagger)
        for token in doc.tokens:
            token.lemmatize(morph_vocab)
        doc_words = []
        doc_words_normal = []
        for i in range(len(doc.tokens)):
            if (doc.tokens[i].pos in ['ADJ', 'NOUN']) and (len(doc.tokens[i].text) > 3) and (
                    doc.tokens[i].lemma.lower() not in russian_stopwords):
                doc_words.append(doc.tokens[i].lemma)
                doc_words_normal.append(doc.tokens[i].text)
        doc2_words = []
        doc2_words_normal = []
        normal_dict = {}
        for i in range(len(doc_words) - 1):
            doc2_words.append(doc_words[i] + ' ' + doc_words[i + 1])
            doc2_words_normal.append(doc_words_normal[i] + ' ' + doc_words_normal[i + 1])
            lemma = doc_words[i] + ' ' + doc_words[i + 1]
            normal = doc_words_normal[i] + ' ' + doc_words_normal[i + 1]
            normal_dict.update({lemma: normal})
        c2 = collections.Counter()
        doc_words.extend(doc2_words)
        for word in doc2_words:
            c2[word] += 1
        print('\nСамые частые слова: ')
        for word in c2.most_common(15):
            word_cloud += normal_dict.get(word[0]) + '\n'
            # print(word[0])
            #print(normal_dict.get(word[0]))
        print(word_cloud)

        return word_cloud
    except Exception as e:
        print(e)
        return 'Error'


def text_from_pdf(doc):
    text = ''
    for current_page in range(len(doc)):
        page = doc.loadPage(current_page)
        text += page.getText("text")
    return text


def text_from_docx(doc):  # читаем текст из .docx
    text = ''
    for paragraph in doc.paragraphs:
        text += ' ' + paragraph.text
    return text


def text_from_doc(path):  # чтение текста из .doc
    text = textract.process(path)
    return text.decode('utf-8')


def process_scan(dir):  # обработка ворда, титульник которого в виде скана, а остальное текстовое
    img_name = 'image1'  # path.splitext(path.basename(r'{}'.format(dir)))[0]
    if dir.name.split('.')[-1] == 'pdf':
        doc = fitz.open(dir)
        text = text_from_pdf(doc)
        feedback_1 = get_feedback(doc, img_name, 'pdf')
        feedback_1 = feedback_1.replace('\n', ' ')
    elif dir.name.split('.')[-1] == 'docx':
        doc = docx.Document(dir)
        text = text_from_docx(doc)
        feedback_1 = get_feedback(doc, img_name, 'docx')
        feedback_1 = feedback_1.replace('\n', ' ')
        print('\n----------------------------------------\n')

    text_original = feedback_1  # весь текст в верхнем регистре
    text_edit = feedback_1.lower()  # весь текст в нижнем регистре
    str_name = text_original[text_edit.find(
        'студ'):text_edit.find('руково')]  #

    dict = {
        'ФИО': get_name_from_feedback1(str_name),
        'Факультет': find_faculty(text_edit),
        'Кафедра': depart_find(text_edit),
        'Направление': find_direction(text_edit),
        'Профиль': find_profile(text_edit),
        'Тема ВКР': find_theme(text_edit),
        'Частотный анализ слов': most_common_word(text)
    }
    # data = make_data(dict)

    # save_in_docx(data, dir, 'scan')
    return dict


def process_text(dir):  # обработка ворда состоящего только из текста
    if dir.name.split('.')[-1] == 'doc':
        text_edit = text_from_doc(dir).splitlines()
        print('это doc')
    elif dir.name.split('.')[-1] == 'docx':
        print('это docx')
        doc = docx.Document(dir)
        text_edit = text_from_docx(doc).splitlines()
    text_original = ' '.join(
        list(filter(None, text_edit)))  # текст в оригинале
    text_edit = text_original.lower()  # весь текст в нижнем регистре
    titul = text_edit[text_edit.find('ульяновский'):text_edit.find(
        'введение')]  # Оставляем только титульник
    str_name = text_original[titul.find(
        'студент'):titul.find('руководитель')]  #

    dict = {
        'ФИО': get_name_from_feedback1(str_name),
        'Факультет': find_faculty(titul),
        'Кафедра': depart_find(text_edit),
        'Направление': find_direction(titul),
        'Профиль': find_profile(titul),
        'Тема ВКР': find_theme(titul),
        'Частотный анализ слов': most_common_word(text_edit)}

    data = make_data(dict)

    # save_in_docx(data, dir, 'text')

    return dict


def make_data(dict):
    data = [
        'ФИО: {}'.format(dict['ФИО']),
        'Факультет: {}'.format(dict['Факультет']),
        'Кафедра: {}'.format(dict['Кафедра']),
        'Направление: {}'.format(dict['Направление']),
        'Профиль: {}'.format(dict['Профиль']),
        'Тема ВКР: {}'.format(dict['Тема ВКР']),
        'Частотный анализ слов:\n{}'.format(dict['Частотный анализ слов'])]
    return data


def main(doc=None):
    print(type(doc))
    print(doc)
    print('Модель построена')

    directory = os.getcwd()
    # folder = directory + '/Выпуск2019/Бакалавры_scans/'

    # стоп слова
    global russian_stopwords
    russian_stopwords = stopwords.words("russian")
    with open('{}/stop_words.txt'.format(directory), 'r',
              encoding="utf-8") as file_handler:  # добавление стоп слов из файла
        for line in file_handler:
            russian_stopwords.extend([line.rstrip()])

    if not text_or_scan(doc):
        data = process_scan(doc)
    else:
        data = process_text(doc)
    return data


if __name__ == '__main__':
    main()
