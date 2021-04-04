# -*- coding: utf-8 -*-
import glob
import win32com.client
import docx
import os, re
from PIL import Image
import io
import pytesseract
import cv2
import matplotlib.pyplot as plt
import collections
import re
from deeppavlov import configs, build_model, train_model
from deeppavlov.core.commands.utils import parse_config
from nltk.corpus import stopwords
from natasha import (
    Segmenter,
    MorphVocab,
    
    NewsEmbedding,
    NewsMorphTagger,
    NewsSyntaxParser,
    NewsNERTagger,
    
    PER,
    NamesExtractor,

    Doc
)

def get_feedback(document, image_name):
    # Получение(со скана) текста
    image2_info = [] # [rId, name]
    for r in document.part.rels.values():
        if isinstance(r._target, docx.parts.image.ImagePart) and ('{}.'.format(image_name) in r._target.partname):
            image2_info.append(r.rId)
            image2_info.append(os.path.basename(r._target.partname))
            
    doc_part = document.part
    image_part = doc_part.related_parts[image2_info[0]]
    a = image_part.blob


    image = Image.open(io.BytesIO(a))
    image.save(image2_info[1])
    k1 = cv2.imread(image2_info[1])
    gray = cv2.cvtColor(k1, cv2.COLOR_BGR2GRAY)
    threshold = cv2.threshold(gray, 127, 255, cv2.THRESH_BINARY)[1]
    imn = image_name+'.jpeg'
    cv2.imwrite(imn, threshold)

    text = pytesseract.image_to_string(Image.open(image2_info[1]), lang='rus')
    return text


def get_name_from_feedback1(feedback):
    name = ''
    err = 0
    res = None
    while True:
        try:
            res = ner_model([feedback])
            break
        except Exception as e:
            err+=1
            if err > 3:
                print(e)
                print('Ошибка в модели!\nСделано 3 попытки!')
                break
            else:                
                continue  
    if res:
        for id, item in enumerate(res[1][0]):
            if item in ['I-PER', 'B-PER']:
                name += ' '+res[0][0][id]    
        print('Ф.И.О: ', name) 
        return name   
    else:
        print('Ф.И.О: Error') 
        return 'None'    

def find_theme(text):
    try:
        pattern = r"(?<=на тему «).*?(?=»)"
        s_l = re.findall(pattern, text)
        if not s_l:
            pattern = r"(?<=на тему: «).*?(?=»)"
            s_l = re.findall(pattern, text)    
        print('Тема:', s_l[0])
        return s_l[0]
    except:
        print('Тема: Error')
        return 'Error'

def find_profile(text):
    try:
        start_text = text.find('профиль')
        text = text[start_text:]
        pattern = r"(?<=«).*?(?=»)"
        s_l = re.findall(pattern, text)    
        print('Профиль:', s_l[0])
        return s_l[0]
    except:
        print('Профиль: Error')
        return 'Error'

def find_direction(text):
    try:
        start_text = text.find('направление')
        text = text[start_text:]
        pattern = r"(?<=«).*?(?=»)"
        s_l = re.findall(pattern, text)    
        print('Направление:', s_l[0])
        return s_l[0]
    except:
        print('Направление: Error')
        return 'Error'

def find_faculty(text):
    try:
        pattern = r"(?<=факультет ).*?(?=кафедра)"
        a = re.findall(pattern, text)    
        a = re.sub(r' {2,}' , ' ', a[0])
        print('Факультет:', a)
        return a
    except:
        print('Факультет: Error')
        return 'Error'

def get_text_from_docx(doc):
    text = ''
    for paragraph in doc.paragraphs:
        text+= ' '+ paragraph.text
    return text 

def save_in_docx(data, dir, mode):
    name = os.path.basename(dir)
    file = '@{}{}'.format(mode, name)
    if os.path.isfile(file):
        #print('Удаляем!!')
        os.remove(file) 
    doc_save = docx.Document()
    for item in data:
        doc_save.add_paragraph(item)
    doc_save.save(file)

def load_docx(folder):
    for root, dirs, files in os.walk(folder): #переводим DOC в DOCX
        for file in files:
            if file.endswith('doc') and not file.startswith('~') and not file.startswith('@'):
                print('Найден DOC файл!!')
                print(file[:-4])
                word = win32com.client.Dispatch("Word.Application")
                word.visible = 0
                wb = word.Documents.Open(folder+file)
                out_file = os.path.abspath('{}.docx'.format(file[:-4]))
                wb.SaveAs2(out_file, FileFormat=16)
                wb.Close()
                word.Quit()               
    paths = []
    for root, dirs, files in os.walk(folder):
        for file in files:
            if file.endswith('docx') and not file.startswith('~') and not file.startswith('@'):
                paths.append(os.path.join(root, file))
    return paths 

def text_or_scan(file): # True если текст
    doc = docx.Document(file)
    text = []
    for p in doc.paragraphs[:10]:
        text.append(p.text)
    text = ' '.join(text).lower()
    if ('ульяновский' in text and 
        'факультет' in text and 
        'кафедра' in text) or (

        'ульян' in text and 
        'факул' in text and 
        'каф' in text):
        #print('Текст')
        return True 
    else:
        #print('Скан')
        return False

def most_common_word(doc):
    word_cloud = '\n'
    text = []
    for p in doc.paragraphs:
        text.append(p.text)
    text = ''.join(text)
    doc = Doc(text)
    doc.segment(segmenter)
    doc.tag_morph(morph_tagger)
    doc.parse_syntax(syntax_parser)
    doc.tag_ner(ner_tagger)
    for token in doc.tokens:
        token.lemmatize(morph_vocab)
    doc_words = []
    for i in range(len(doc.tokens)):
        if (doc.tokens[i].pos in ['ADJ', 'NOUN']) and (len(doc.tokens[i].text) > 3) and (doc.tokens[i].lemma.lower() not in russian_stopwords):
            doc_words.append(doc.tokens[i].lemma)
    doc2_words = []
    for i in range(len(doc_words) - 1):
        doc2_words.append(doc_words[i] + ' ' + doc_words[i + 1])
    c2 = collections.Counter()
    doc_words.extend(doc2_words)
    for word in doc2_words:
        c2[word] += 1
    print('\nСамые частые слова: ')
    for word in c2.most_common(15):
        word_cloud += word[0] + '\n'
        print(word[0])
    return word_cloud
def process_scan(dir):
    data = []
    doc = docx.Document(dir) 
    feedback_1 = get_feedback(doc, 'image1')
    feedback_1 = feedback_1.replace('\n', ' ')
    print('\n----------------------------------------\n')
    text_original = feedback_1 #весь текст в верхнем регистре            
    text_edit = feedback_1.lower() #весь текст в нижнем регистре
    str_name = text_original[text_edit.find('студ'):text_edit.find('руково')]#  

    data.append('ФИО: {}'.format(get_name_from_feedback1(str_name)))
    data.append('Факультет: {}'.format(find_faculty(text_edit)))
    data.append('Направление: {}'.format(find_direction(text_edit)))
    data.append('Профиль: {}'.format(find_profile(text_edit)))
    data.append('Тема ВКР: {}'.format(find_theme(text_edit)))    
    data.append('Частотный анализ слов:\n{}'.format(most_common_word(doc)))
    save_in_docx(data, dir, 'scan')
def process_text(dir):  
    data = []
    doc = docx.Document(dir)
    text_edit = get_text_from_docx(doc).splitlines()
    text_original = ' '.join(list(filter(None, text_edit))) #весь текст в верхнем регистре
    text_edit = text_original.lower() #весь текст в нижнем регистре
    titul = text_edit[text_edit.find('ульяновский'):text_edit.find('введение')] #Оставляем только титульник
    str_name = text_original[titul.find('студент'):titul.find('руководитель')]#
    data.append('ФИО: {}'.format(get_name_from_feedback1(str_name))) #
    data.append('Факультет: {}'.format(find_faculty(titul)))
    data.append('Направление: {}'.format(find_direction(titul)))
    data.append('Профиль: {}'.format(find_profile(titul)))
    data.append('Тема ВКР: {}'.format(find_theme(titul)))    
    data.append('Частотный анализ слов:\n{}'.format(most_common_word(doc)))
    save_in_docx(data, dir, 'text') 
if __name__ == '__main__':
    segmenter = Segmenter()
    morph_vocab = MorphVocab()
    emb = NewsEmbedding()
    morph_tagger = NewsMorphTagger(emb)
    syntax_parser = NewsSyntaxParser(emb)
    ner_tagger = NewsNERTagger(emb)
    names_extractor = NamesExtractor(morph_vocab)
    
    ner_model = build_model(configs.ner.ner_rus_bert, download=False)
    print('Модель построена')
    pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    plt.rcParams['figure.figsize'] = (30, 15)
    directory = os.getcwd()
    folder = directory + '/Выпуск2019/Бакалавры_texts/'
    
    #стоп слова
    russian_stopwords = stopwords.words("russian")
    with open('{}/stop_words.txt'.format(directory), 'r', encoding="utf-8") as file_handler:
        for line in file_handler:
            russian_stopwords.extend([line.rstrip()])
    print(russian_stopwords)
    os.chdir(folder)
    paths = load_docx(folder)
    print(paths)
    for dir in paths: #пробегаемся по всем word файлам в папке folder
        print(dir) 
        if not text_or_scan(dir):
            print('Это скан!!!')
            process_scan(dir)       
        else:
            print('Это текст!!')
            process_text(dir)
            print('\n----------------------------------------\n')

            

