import json

import docx
from natasha import (
    Segmenter,

    NewsEmbedding,
    NewsMorphTagger,
    Doc, MorphVocab
)


def load_stopwords():
    with open('stopwords.json', encoding='utf-8') as json_file:
        data = json.load(json_file)
        stop_arr = data['stopwords']
    return stop_arr


# TODO: Добавить функцию пополнения стоплиста
#
# def getText(filename):
#     doc = Document(filename)
#     fullText = []
#     for para in doc.paragraphs:
#         if fullText:
#             fullText[-1] += para.text
#         else:
#             fullText.append('')
#     return '\n'.join(fullText)
#
# a = getText("./Валитов-ВКР.docx")
# for i in a:
#     print(i)
#


def list(stop_arr):
    emb = NewsEmbedding()
    # TODO: Вынести это в отдельную функцию
    document = docx.Document('Валитов-ВКР.docx')
    text = '\n'.join([para.text for para in document.paragraphs])
    segmenter = Segmenter()
    doc = Doc(text)
    doc.segment(segmenter)

    morph_tagger = NewsMorphTagger(emb)
    doc.tag_morph(morph_tagger)

    morph_vocab = MorphVocab()
    for token in doc.tokens:
        token.lemmatize(morph_vocab)
    for token in doc.tokens:
        if token.lemma not in stop_arr and not token.lemma.isnumeric():
            yield token.lemma.lower()


stopArr = load_stopwords()

dict = dict()
for i in list(stopArr):
    try:
        dict[i] += 1
    except:
        dict[i] = 1

print({k: v for k, v in reversed(sorted(dict.items(), key=lambda item: item[1]))})
